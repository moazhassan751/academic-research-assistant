"""
Export utilities for Academic Research Assistant
Supports multiple output formats: PDF, Word, LaTeX, CSV, HTML
"""

import os
import csv
import json
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import logging

try:
    # PDF generation
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    # Word document generation
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    # HTML to PDF conversion (alternative)
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportManager:
    """Manages export functionality for research outputs"""
    
    def __init__(self):
        self.supported_formats = {
            'markdown': True,
            'txt': True,
            'json': True,
            'csv': True,
            'html': True,
            'latex': True,
            'pdf': PDF_AVAILABLE or PDFKIT_AVAILABLE,
            'docx': DOCX_AVAILABLE
        }
        
    def get_supported_formats(self) -> Dict[str, bool]:
        """Return dict of supported formats and their availability"""
        return self.supported_formats.copy()
    
    def export_draft(self, draft: Dict[str, Any], output_path: str, 
                    format_type: str = 'markdown') -> bool:
        """
        Export research draft to specified format
        
        Args:
            draft: Draft content dictionary
            output_path: Output file path (without extension)
            format_type: Export format ('markdown', 'pdf', 'docx', 'latex', 'html')
        
        Returns:
            bool: Success status
        """
        try:
            format_type = format_type.lower()
            
            if format_type not in self.supported_formats:
                raise ValueError(f"Unsupported format: {format_type}")
            
            if not self.supported_formats[format_type]:
                raise ValueError(f"Format {format_type} is not available (missing dependencies)")
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            if format_type == 'markdown':
                return self._export_draft_markdown(draft, f"{output_path}.md")
            elif format_type == 'pdf':
                return self._export_draft_pdf(draft, f"{output_path}.pdf")
            elif format_type == 'docx':
                return self._export_draft_docx(draft, f"{output_path}.docx")
            elif format_type == 'latex':
                return self._export_draft_latex(draft, f"{output_path}.tex")
            elif format_type == 'html':
                return self._export_draft_html(draft, f"{output_path}.html")
            elif format_type == 'json':
                return self._export_draft_json(draft, f"{output_path}.json")
            else:
                raise ValueError(f"Format {format_type} handler not implemented")
                
        except Exception as e:
            logger.error(f"Error exporting draft to {format_type}: {e}")
            return False
    
    def export_bibliography(self, bibliography: str, papers: List[Any], 
                          output_path: str, format_type: str = 'txt') -> bool:
        """
        Export bibliography to specified format
        
        Args:
            bibliography: Bibliography text content
            papers: List of paper objects
            output_path: Output file path (without extension)
            format_type: Export format ('txt', 'csv', 'bibtex', 'json', 'pdf', 'docx')
        
        Returns:
            bool: Success status
        """
        try:
            format_type = format_type.lower()
            
            # Map bibtex to latex for consistency
            if format_type == 'bibtex':
                format_type = 'latex'
            
            if format_type not in self.supported_formats:
                raise ValueError(f"Unsupported format: {format_type}")
            
            if not self.supported_formats[format_type]:
                raise ValueError(f"Format {format_type} is not available (missing dependencies)")
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            if format_type == 'txt':
                return self._export_bibliography_txt(bibliography, f"{output_path}.txt")
            elif format_type == 'csv':
                return self._export_bibliography_csv(papers, f"{output_path}.csv")
            elif format_type == 'json':
                return self._export_bibliography_json(papers, f"{output_path}.json")
            elif format_type == 'latex':
                return self._export_bibliography_bibtex(papers, f"{output_path}.bib")
            elif format_type == 'pdf':
                return self._export_bibliography_pdf(bibliography, f"{output_path}.pdf")
            elif format_type == 'docx':
                return self._export_bibliography_docx(bibliography, f"{output_path}.docx")
            elif format_type == 'markdown':
                return self._export_bibliography_markdown(bibliography, papers, f"{output_path}.md")
            elif format_type == 'html':
                return self._export_bibliography_html(bibliography, papers, f"{output_path}.html")
            else:
                raise ValueError(f"Format {format_type} handler not implemented")
                
        except Exception as e:
            logger.error(f"Error exporting bibliography to {format_type}: {e}")
            return False
    
    # Draft export methods
    def _export_draft_markdown(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as Markdown"""
        try:
            md_content = self._format_draft_as_markdown(draft)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(md_content)
            logger.info(f"Draft exported to Markdown: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting draft to Markdown: {e}")
            return False
    
    def _export_draft_pdf(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as PDF"""
        try:
            if PDF_AVAILABLE:
                return self._export_draft_pdf_reportlab(draft, filepath)
            elif PDFKIT_AVAILABLE:
                return self._export_draft_pdf_pdfkit(draft, filepath)
            else:
                raise ValueError("No PDF library available")
        except Exception as e:
            logger.error(f"Error exporting draft to PDF: {e}")
            return False
    
    def _export_draft_pdf_reportlab(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as PDF using ReportLab"""
        try:
            doc = SimpleDocTemplate(filepath, pagesize=A4,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                alignment=TA_CENTER,
                spaceAfter=30
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=20,
                spaceAfter=12
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceBefore=6,
                spaceAfter=6
            )
            
            story = []
            
            # Title
            title = draft.get('title', 'Research Paper Draft')
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            # Abstract
            if 'abstract' in draft and draft['abstract']:
                story.append(Paragraph("Abstract", heading_style))
                story.append(Paragraph(str(draft['abstract']), normal_style))
                story.append(Spacer(1, 12))
            
            # Introduction
            if 'introduction' in draft and draft['introduction']:
                story.append(Paragraph("1. Introduction", heading_style))
                story.append(Paragraph(str(draft['introduction']), normal_style))
                story.append(Spacer(1, 12))
            
            # Sections
            section_num = 2
            if 'sections' in draft and draft['sections']:
                for key, section in draft['sections'].items():
                    if key.startswith('theme_') and isinstance(section, dict):
                        title = section.get('title', f'Section {section_num}')
                        content = section.get('content', 'Content unavailable')
                        story.append(Paragraph(f"{section_num}. {title}", heading_style))
                        story.append(Paragraph(str(content), normal_style))
                        story.append(Spacer(1, 12))
                        section_num += 1
            
            # Discussion
            if 'discussion' in draft and draft['discussion']:
                story.append(Paragraph(f"{section_num}. Discussion", heading_style))
                story.append(Paragraph(str(draft['discussion']), normal_style))
                story.append(Spacer(1, 12))
                section_num += 1
            
            # Conclusion
            if 'conclusion' in draft and draft['conclusion']:
                story.append(Paragraph(f"{section_num}. Conclusion", heading_style))
                story.append(Paragraph(str(draft['conclusion']), normal_style))
                story.append(Spacer(1, 12))
            
            # References
            if 'bibliography' in draft and draft['bibliography']:
                story.append(PageBreak())
                story.append(Paragraph("References", heading_style))
                story.append(Paragraph(str(draft['bibliography']), normal_style))
            
            doc.build(story)
            logger.info(f"Draft exported to PDF: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating PDF with ReportLab: {e}")
            return False
    
    def _export_draft_pdf_pdfkit(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as PDF using pdfkit (HTML to PDF)"""
        try:
            # First create HTML
            html_content = self._format_draft_as_html(draft)
            
            # Convert to PDF
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None
            }
            
            pdfkit.from_string(html_content, filepath, options=options)
            logger.info(f"Draft exported to PDF: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating PDF with pdfkit: {e}")
            return False
    
    def _export_draft_docx(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as Word document"""
        try:
            doc = Document()
            
            # Title
            title = draft.get('title', 'Research Paper Draft')
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Abstract
            if 'abstract' in draft and draft['abstract']:
                doc.add_heading('Abstract', level=1)
                doc.add_paragraph(str(draft['abstract']))
            
            # Introduction
            if 'introduction' in draft and draft['introduction']:
                doc.add_heading('1. Introduction', level=1)
                doc.add_paragraph(str(draft['introduction']))
            
            # Sections
            section_num = 2
            if 'sections' in draft and draft['sections']:
                for key, section in draft['sections'].items():
                    if key.startswith('theme_') and isinstance(section, dict):
                        title = section.get('title', f'Section {section_num}')
                        content = section.get('content', 'Content unavailable')
                        doc.add_heading(f'{section_num}. {title}', level=1)
                        doc.add_paragraph(str(content))
                        section_num += 1
            
            # Discussion
            if 'discussion' in draft and draft['discussion']:
                doc.add_heading(f'{section_num}. Discussion', level=1)
                doc.add_paragraph(str(draft['discussion']))
                section_num += 1
            
            # Conclusion
            if 'conclusion' in draft and draft['conclusion']:
                doc.add_heading(f'{section_num}. Conclusion', level=1)
                doc.add_paragraph(str(draft['conclusion']))
            
            # References
            if 'bibliography' in draft and draft['bibliography']:
                doc.add_page_break()
                doc.add_heading('References', level=1)
                doc.add_paragraph(str(draft['bibliography']))
            
            doc.save(filepath)
            logger.info(f"Draft exported to Word: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return False
    
    def _export_draft_latex(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as LaTeX document"""
        try:
            latex_content = self._format_draft_as_latex(draft)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(latex_content)
            logger.info(f"Draft exported to LaTeX: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting draft to LaTeX: {e}")
            return False
    
    def _export_draft_html(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as HTML document"""
        try:
            html_content = self._format_draft_as_html(draft)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)
            logger.info(f"Draft exported to HTML: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting draft to HTML: {e}")
            return False
    
    def _export_draft_json(self, draft: Dict[str, Any], filepath: str) -> bool:
        """Export draft as JSON document"""
        try:
            import json
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(draft, f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Draft exported to JSON: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting draft to JSON: {e}")
            return False
    
    # Bibliography export methods
    def _export_bibliography_txt(self, bibliography: str, filepath: str) -> bool:
        """Export bibliography as plain text"""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(bibliography)
            logger.info(f"Bibliography exported to text: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting bibliography to text: {e}")
            return False
    
    def _export_bibliography_csv(self, papers: List[Any], filepath: str) -> bool:
        """Export bibliography as CSV"""
        try:
            with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['title', 'authors', 'year', 'journal', 'doi', 'url', 'abstract']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                
                writer.writeheader()
                for paper in papers:
                    writer.writerow({
                        'title': getattr(paper, 'title', ''),
                        'authors': '; '.join(getattr(paper, 'authors', [])),
                        'year': getattr(paper.published_date, 'year', '') if hasattr(paper, 'published_date') and paper.published_date else '',
                        'journal': getattr(paper, 'journal', ''),
                        'doi': getattr(paper, 'doi', ''),
                        'url': getattr(paper, 'url', ''),
                        'abstract': getattr(paper, 'abstract', '')[:500] + '...' if len(getattr(paper, 'abstract', '')) > 500 else getattr(paper, 'abstract', '')
                    })
            
            logger.info(f"Bibliography exported to CSV: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting bibliography to CSV: {e}")
            return False
    
    def _export_bibliography_json(self, papers: List[Any], filepath: str) -> bool:
        """Export bibliography as JSON"""
        try:
            papers_data = []
            for paper in papers:
                paper_dict = {
                    'title': getattr(paper, 'title', ''),
                    'authors': getattr(paper, 'authors', []),
                    'year': getattr(paper.published_date, 'year', None) if hasattr(paper, 'published_date') and paper.published_date else None,
                    'journal': getattr(paper, 'journal', ''),
                    'doi': getattr(paper, 'doi', ''),
                    'url': getattr(paper, 'url', ''),
                    'abstract': getattr(paper, 'abstract', ''),
                    'keywords': getattr(paper, 'keywords', [])
                }
                papers_data.append(paper_dict)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(papers_data, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info(f"Bibliography exported to JSON: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting bibliography to JSON: {e}")
            return False
    
    def _export_bibliography_bibtex(self, papers: List[Any], filepath: str) -> bool:
        """Export bibliography as BibTeX"""
        try:
            bibtex_content = self._format_papers_as_bibtex(papers)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(bibtex_content)
            logger.info(f"Bibliography exported to BibTeX: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error exporting bibliography to BibTeX: {e}")
            return False
    
    def _export_bibliography_pdf(self, bibliography: str, filepath: str) -> bool:
        """Export bibliography as PDF"""
        try:
            if PDF_AVAILABLE:
                doc = SimpleDocTemplate(filepath, pagesize=A4,
                                      rightMargin=72, leftMargin=72,
                                      topMargin=72, bottomMargin=18)
                
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'BibTitle',
                    parent=styles['Heading1'],
                    fontSize=16,
                    alignment=TA_CENTER,
                    spaceAfter=30
                )
                normal_style = ParagraphStyle(
                    'BibNormal',
                    parent=styles['Normal'],
                    fontSize=10,
                    spaceBefore=6,
                    spaceAfter=6
                )
                
                story = []
                story.append(Paragraph("Bibliography", title_style))
                story.append(Spacer(1, 12))
                
                # Split bibliography into entries and format each
                entries = bibliography.split('\n\n')
                for entry in entries:
                    if entry.strip():
                        story.append(Paragraph(entry.strip(), normal_style))
                        story.append(Spacer(1, 6))
                
                doc.build(story)
                logger.info(f"Bibliography exported to PDF: {filepath}")
                return True
            else:
                raise ValueError("PDF library not available")
                
        except Exception as e:
            logger.error(f"Error exporting bibliography to PDF: {e}")
            return False
    
    def _export_bibliography_docx(self, bibliography: str, filepath: str) -> bool:
        """Export bibliography as Word document"""
        try:
            doc = Document()
            
            # Title
            title_para = doc.add_heading('Bibliography', level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Bibliography entries
            entries = bibliography.split('\n\n')
            for entry in entries:
                if entry.strip():
                    doc.add_paragraph(entry.strip())
            
            doc.save(filepath)
            logger.info(f"Bibliography exported to Word: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting bibliography to Word: {e}")
            return False
    
    def _export_bibliography_markdown(self, bibliography: str, papers: List[Any], filepath: str) -> bool:
        """Export bibliography as Markdown"""
        try:
            md_content = "# Bibliography\n\n"
            md_content += bibliography + "\n\n"
            
            # Add detailed paper information
            if papers:
                md_content += "## Detailed References\n\n"
                for i, paper in enumerate(papers, 1):
                    md_content += f"### {i}. {paper.title}\n\n"
                    
                    if paper.authors:
                        authors_str = ", ".join(paper.authors)
                        md_content += f"**Authors:** {authors_str}\n\n"
                    
                    year = getattr(paper, 'year', None)
                    if year:
                        md_content += f"**Year:** {year}\n\n"
                    
                    if hasattr(paper, 'journal') and paper.journal:
                        md_content += f"**Journal:** {paper.journal}\n\n"
                    
                    if hasattr(paper, 'doi') and paper.doi:
                        md_content += f"**DOI:** {paper.doi}\n\n"
                    
                    if hasattr(paper, 'url') and paper.url:
                        md_content += f"**URL:** [{paper.url}]({paper.url})\n\n"
                    
                    if hasattr(paper, 'abstract') and paper.abstract:
                        md_content += f"**Abstract:** {paper.abstract}\n\n"
                    
                    md_content += "---\n\n"
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            logger.info(f"Bibliography exported to Markdown: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting bibliography to Markdown: {e}")
            return False
    
    def _export_bibliography_html(self, bibliography: str, papers: List[Any], filepath: str) -> bool:
        """Export bibliography as HTML"""
        try:
            html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Bibliography</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }
        h1 { color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }
        h2 { color: #555; margin-top: 30px; }
        h3 { color: #666; margin-top: 25px; }
        .paper { margin-bottom: 20px; padding: 15px; border-left: 4px solid #007acc; background-color: #f9f9f9; }
        .authors { font-weight: bold; color: #444; }
        .year { font-style: italic; color: #666; }
        .journal { color: #555; }
        .doi { color: #007acc; }
        .abstract { margin-top: 10px; color: #333; }
        a { color: #007acc; text-decoration: none; }
        a:hover { text-decoration: underline; }
        hr { border: none; border-top: 1px solid #ddd; margin: 20px 0; }
    </style>
</head>
<body>
    <h1>Bibliography</h1>
"""
            
            # Add basic bibliography content
            bibliography_lines = bibliography.split('\n')
            html_content += "    <div class='summary'>\n"
            for line in bibliography_lines:
                if line.strip():
                    html_content += f"        <p>{line}</p>\n"
            html_content += "    </div>\n\n"
            
            # Add detailed paper information
            if papers:
                html_content += "    <h2>Detailed References</h2>\n\n"
                for i, paper in enumerate(papers, 1):
                    html_content += f"    <div class='paper'>\n"
                    html_content += f"        <h3>{i}. {paper.title}</h3>\n"
                    
                    if paper.authors:
                        authors_str = ", ".join(paper.authors)
                        html_content += f"        <p class='authors'>Authors: {authors_str}</p>\n"
                    
                    year = getattr(paper, 'year', None)
                    if year:
                        html_content += f"        <p class='year'>Year: {year}</p>\n"
                    
                    if hasattr(paper, 'journal') and paper.journal:
                        html_content += f"        <p class='journal'>Journal: {paper.journal}</p>\n"
                    
                    if hasattr(paper, 'doi') and paper.doi:
                        html_content += f"        <p class='doi'>DOI: {paper.doi}</p>\n"
                    
                    if hasattr(paper, 'url') and paper.url:
                        html_content += f"        <p>URL: <a href='{paper.url}' target='_blank'>{paper.url}</a></p>\n"
                    
                    if hasattr(paper, 'abstract') and paper.abstract:
                        abstract_preview = paper.abstract[:500] + "..." if len(paper.abstract) > 500 else paper.abstract
                        html_content += f"        <p class='abstract'><strong>Abstract:</strong> {abstract_preview}</p>\n"
                    
                    html_content += "    </div>\n\n"
            
            html_content += """</body>
</html>"""
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"Bibliography exported to HTML: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting bibliography to HTML: {e}")
            return False

    # Formatting helper methods
    def _format_draft_as_markdown(self, draft: Dict[str, Any]) -> str:
        """Format draft as professionally structured markdown document"""
        md_content = ""
        
        # Title with proper formatting
        title = draft.get('title', 'Research Paper Draft')
        md_content += f"# {title}\n\n"
        
        # Add metadata section
        md_content += "---\n"
        md_content += f"**Document Type:** Research Paper Draft\n"
        md_content += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        md_content += f"**Status:** Draft\n"
        md_content += "---\n\n"
        
        # Table of Contents
        md_content += "## Table of Contents\n\n"
        toc_items = []
        if 'abstract' in draft and draft['abstract']:
            toc_items.append("- [Abstract](#abstract)")
        if 'introduction' in draft and draft['introduction']:
            toc_items.append("- [1. Introduction](#1-introduction)")
        
        section_num = 2
        if 'sections' in draft and draft['sections']:
            for key, section in draft['sections'].items():
                if key.startswith('theme_') and isinstance(section, dict):
                    title = section.get('title', f'Section {section_num}')
                    clean_title = title.lower().replace(' ', '-').replace(',', '').replace('.', '')
                    toc_items.append(f"- [{section_num}. {title}](#{section_num}-{clean_title})")
                    section_num += 1
        
        if 'discussion' in draft and draft['discussion']:
            toc_items.append(f"- [{section_num}. Discussion](#{section_num}-discussion)")
            section_num += 1
        if 'conclusion' in draft and draft['conclusion']:
            toc_items.append(f"- [{section_num}. Conclusion](#{section_num}-conclusion)")
        if 'bibliography' in draft and draft['bibliography']:
            toc_items.append("- [References](#references)")
        
        md_content += "\n".join(toc_items) + "\n\n"
        
        # Abstract with enhanced formatting
        if 'abstract' in draft and draft['abstract']:
            md_content += "## Abstract\n\n"
            md_content += "> **Summary:** This section provides a concise overview of the research findings and methodology.\n\n"
            abstract_text = self._format_content_with_structure(str(draft['abstract']))
            md_content += abstract_text + "\n\n"
            md_content += "---\n\n"
        
        # Introduction with enhanced formatting
        if 'introduction' in draft and draft['introduction']:
            md_content += "## 1. Introduction\n\n"
            md_content += "### 1.1 Background\n\n"
            intro_text = self._format_content_with_structure(str(draft['introduction']))
            md_content += intro_text + "\n\n"
        
        # Main sections with improved structure
        section_num = 2
        if 'sections' in draft and draft['sections']:
            for key, section in draft['sections'].items():
                if key.startswith('theme_') and isinstance(section, dict):
                    title = section.get('title', f'Section {section_num}')
                    content = section.get('content', 'Content unavailable')
                    
                    md_content += f"## {section_num}. {title}\n\n"
                    
                    # Add subsection structure
                    md_content += f"### {section_num}.1 Overview\n\n"
                    formatted_content = self._format_content_with_structure(str(content))
                    md_content += formatted_content + "\n\n"
                    
                    # Add key findings if content is substantial
                    if len(str(content)) > 200:
                        md_content += f"### {section_num}.2 Key Insights\n\n"
                        insights = self._extract_key_points(str(content))
                        md_content += insights + "\n\n"
                    
                    section_num += 1
        
        # Discussion with enhanced structure
        if 'discussion' in draft and draft['discussion']:
            md_content += f"## {section_num}. Discussion\n\n"
            md_content += f"### {section_num}.1 Analysis\n\n"
            discussion_text = self._format_content_with_structure(str(draft['discussion']))
            md_content += discussion_text + "\n\n"
            
            md_content += f"### {section_num}.2 Implications\n\n"
            md_content += "The findings presented in this research have several important implications:\n\n"
            md_content += "- **Theoretical Implications:** Further research in this area\n"
            md_content += "- **Practical Applications:** Real-world implementation considerations\n"
            md_content += "- **Future Directions:** Areas for continued investigation\n\n"
            section_num += 1
        
        # Conclusion with enhanced structure
        if 'conclusion' in draft and draft['conclusion']:
            md_content += f"## {section_num}. Conclusion\n\n"
            md_content += "### Summary of Findings\n\n"
            conclusion_text = self._format_content_with_structure(str(draft['conclusion']))
            md_content += conclusion_text + "\n\n"
            
            md_content += "### Final Remarks\n\n"
            md_content += "This research contributes to the ongoing understanding of the topic and provides a foundation for future investigations.\n\n"
        
        # References with proper formatting
        if 'bibliography' in draft and draft['bibliography']:
            md_content += "## References\n\n"
            md_content += "---\n\n"
            bib_text = str(draft['bibliography'])
            if isinstance(draft['bibliography'], list):
                for i, ref in enumerate(draft['bibliography'], 1):
                    md_content += f"{i}. {ref}\n\n"
            else:
                md_content += bib_text + "\n\n"
        
        # Add appendices section if there's additional data
        if 'data' in draft or 'appendices' in draft:
            md_content += "## Appendices\n\n"
            md_content += "### Appendix A: Additional Data\n\n"
            md_content += "Supplementary materials and data are available upon request.\n\n"
        
        return md_content
    
    def _format_content_with_structure(self, content: str) -> str:
        """Format content with proper paragraph structure and bullet points"""
        if not content or len(content.strip()) == 0:
            return "Content not available."
        
        # Split content into sentences
        sentences = [s.strip() for s in content.split('.') if s.strip()]
        
        if len(sentences) <= 2:
            return content
        
        formatted_content = ""
        
        # Group sentences into paragraphs
        current_paragraph = []
        for i, sentence in enumerate(sentences):
            current_paragraph.append(sentence)
            
            # Create paragraph breaks every 2-3 sentences or at logical breaks
            if (len(current_paragraph) >= 3 or 
                i == len(sentences) - 1 or
                any(keyword in sentence.lower() for keyword in ['however', 'therefore', 'furthermore', 'additionally', 'moreover'])):
                
                paragraph_text = '. '.join(current_paragraph)
                if not paragraph_text.endswith('.'):
                    paragraph_text += '.'
                formatted_content += paragraph_text + "\n\n"
                current_paragraph = []
        
        # Look for lists or enumerable items
        if any(keyword in content.lower() for keyword in ['first', 'second', 'third', 'include', 'such as', 'following']):
            formatted_content += self._create_bullet_points(content)
        
        return formatted_content.strip()
    
    def _extract_key_points(self, content: str) -> str:
        """Extract key points and format as bullet points"""
        sentences = [s.strip() for s in content.split('.') if s.strip() and len(s.strip()) > 20]
        
        if len(sentences) < 2:
            return "- " + content
        
        key_points = []
        for sentence in sentences[:4]:  # Take up to 4 key sentences
            if len(sentence) > 30:  # Only substantial sentences
                # Clean up the sentence
                clean_sentence = sentence.strip()
                if not clean_sentence.endswith('.'):
                    clean_sentence += '.'
                key_points.append(f"- **{clean_sentence}**")
        
        return "\n".join(key_points) if key_points else f"- {content}"
    
    def _create_bullet_points(self, content: str) -> str:
        """Create structured bullet points from content"""
        bullet_content = "\n#### Key Points:\n\n"
        
        # Look for natural list items
        if 'include' in content.lower() or 'such as' in content.lower():
            items = content.split(',')
            if len(items) > 1:
                for item in items[:5]:  # Limit to 5 items
                    clean_item = item.strip().replace('include', '').replace('such as', '').strip()
                    if len(clean_item) > 5:
                        bullet_content += f"- {clean_item}\n"
        else:
            # Create thematic bullet points
            sentences = [s.strip() for s in content.split('.') if s.strip() and len(s.strip()) > 20]
            for sentence in sentences[:3]:  # Limit to 3 points
                bullet_content += f"- {sentence}.\n"
        
        return bullet_content + "\n"
    
    def _format_draft_as_latex(self, draft: Dict[str, Any]) -> str:
        """Format draft as LaTeX document"""
        latex_content = """\\documentclass[12pt,a4paper]{article}
\\usepackage[utf8]{inputenc}
\\usepackage[english]{babel}
\\usepackage{amsmath}
\\usepackage{amsfonts}
\\usepackage{amssymb}
\\usepackage{graphicx}
\\usepackage[left=2cm,right=2cm,top=2cm,bottom=2cm]{geometry}
\\usepackage{cite}
\\usepackage{url}

\\title{""" + draft.get('title', 'Research Paper Draft').replace('&', '\\&').replace('_', '\\_') + """}
\\author{Academic Research Assistant}
\\date{\\today}

\\begin{document}

\\maketitle

"""
        
        # Abstract
        if 'abstract' in draft and draft['abstract']:
            latex_content += "\\begin{abstract}\n"
            latex_content += self._escape_latex(str(draft['abstract'])) + "\n"
            latex_content += "\\end{abstract}\n\n"
        
        # Introduction
        if 'introduction' in draft and draft['introduction']:
            latex_content += "\\section{Introduction}\n"
            latex_content += self._escape_latex(str(draft['introduction'])) + "\n\n"
        
        # Sections
        if 'sections' in draft and draft['sections']:
            for key, section in draft['sections'].items():
                if key.startswith('theme_') and isinstance(section, dict):
                    title = section.get('title', 'Section')
                    content = section.get('content', 'Content unavailable')
                    latex_content += f"\\section{{{self._escape_latex(title)}}}\n"
                    latex_content += self._escape_latex(str(content)) + "\n\n"
        
        # Discussion
        if 'discussion' in draft and draft['discussion']:
            latex_content += "\\section{Discussion}\n"
            latex_content += self._escape_latex(str(draft['discussion'])) + "\n\n"
        
        # Conclusion
        if 'conclusion' in draft and draft['conclusion']:
            latex_content += "\\section{Conclusion}\n"
            latex_content += self._escape_latex(str(draft['conclusion'])) + "\n\n"
        
        # References
        if 'bibliography' in draft and draft['bibliography']:
            latex_content += "\\section*{References}\n"
            # Simple bibliography formatting
            bib_entries = str(draft['bibliography']).split('\n\n')
            for entry in bib_entries:
                if entry.strip():
                    latex_content += self._escape_latex(entry.strip()) + "\n\n"
        
        latex_content += "\\end{document}"
        return latex_content
    
    def _format_draft_as_html(self, draft: Dict[str, Any]) -> str:
        """Format draft as professionally structured HTML document"""
        title = draft.get('title', 'Research Paper Draft')
        
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self._escape_html(title)}</title>
    <style>
        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            line-height: 1.8;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px;
            background-color: #fff;
            color: #333;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 40px;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 20px;
        }}
        
        h1 {{
            color: #2c3e50;
            font-size: 2.5em;
            margin-bottom: 10px;
            font-weight: 300;
        }}
        
        .metadata {{
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            padding: 15px;
            border-radius: 8px;
            margin: 20px 0;
            font-size: 0.9em;
            color: #666;
        }}
        
        .toc {{
            background-color: #f8f9fa;
            border: 1px solid #dee2e6;
            border-radius: 8px;
            padding: 20px;
            margin: 30px 0;
        }}
        
        .toc h2 {{
            color: #495057;
            font-size: 1.3em;
            margin-bottom: 15px;
        }}
        
        .toc ul {{
            list-style: none;
            padding-left: 0;
        }}
        
        .toc li {{
            padding: 5px 0;
            border-bottom: 1px dotted #ccc;
        }}
        
        .toc a {{
            color: #007bff;
            text-decoration: none;
            font-weight: 500;
        }}
        
        .toc a:hover {{
            text-decoration: underline;
        }}
        
        h2 {{
            color: #2c3e50;
            font-size: 1.8em;
            margin-top: 40px;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid #3498db;
        }}
        
        h3 {{
            color: #34495e;
            font-size: 1.4em;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        
        h4 {{
            color: #7f8c8d;
            font-size: 1.2em;
            margin-top: 25px;
            margin-bottom: 10px;
        }}
        
        p {{
            text-align: justify;
            margin-bottom: 20px;
            font-size: 1.1em;
            text-indent: 1.5em;
        }}
        
        .abstract {{
            background: linear-gradient(135deg, #e8f4fd 0%, #c3e8ff 100%);
            padding: 25px;
            border-left: 5px solid #3498db;
            margin: 30px 0;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        
        .abstract h2 {{
            margin-top: 0;
            color: #2980b9;
        }}
        
        .abstract p {{
            font-style: italic;
            text-indent: 0;
            color: #2c3e50;
        }}
        
        .key-points {{
            background-color: #fff3cd;
            border: 1px solid #ffeaa7;
            border-radius: 8px;
            padding: 20px;
            margin: 25px 0;
        }}
        
        .key-points h4 {{
            color: #856404;
            margin-top: 0;
        }}
        
        ul {{
            padding-left: 25px;
        }}
        
        li {{
            margin-bottom: 8px;
            line-height: 1.6;
        }}
        
        .strong-point {{
            font-weight: 600;
            color: #2c3e50;
        }}
        
        .section-overview {{
            background-color: #f8f9fa;
            border-left: 4px solid #6c757d;
            padding: 15px;
            margin: 20px 0;
            font-style: italic;
        }}
        
        .implications {{
            background-color: #e8f5e8;
            border: 1px solid #c3e6c3;
            border-radius: 8px;
            padding: 20px;
            margin: 25px 0;
        }}
        
        .references {{
            background-color: #f8f9fa;
            padding: 25px;
            border-radius: 8px;
            margin-top: 40px;
            font-size: 0.95em;
            line-height: 1.6;
        }}
        
        .references h2 {{
            color: #495057;
        }}
        
        .footer {{
            text-align: center;
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #dee2e6;
            color: #6c757d;
            font-size: 0.9em;
        }}
        
        @media print {{
            body {{
                font-size: 12pt;
                line-height: 1.6;
            }}
            .toc {{
                page-break-after: always;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{self._escape_html(title)}</h1>
        <div class="metadata">
            <strong>Document Type:</strong> Research Paper Draft<br>
            <strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br>
            <strong>Status:</strong> Draft
        </div>
    </div>
"""
        
        # Table of Contents
        html_content += '    <div class="toc">\n'
        html_content += '        <h2>Table of Contents</h2>\n'
        html_content += '        <ul>\n'
        
        if 'abstract' in draft and draft['abstract']:
            html_content += '            <li><a href="#abstract">Abstract</a></li>\n'
        if 'introduction' in draft and draft['introduction']:
            html_content += '            <li><a href="#introduction">1. Introduction</a></li>\n'
        
        section_num = 2
        if 'sections' in draft and draft['sections']:
            for key, section in draft['sections'].items():
                if key.startswith('theme_') and isinstance(section, dict):
                    title_clean = section.get('title', f'Section {section_num}')
                    anchor = title_clean.lower().replace(' ', '-').replace(',', '').replace('.', '')
                    html_content += f'            <li><a href="#section-{section_num}">{section_num}. {self._escape_html(title_clean)}</a></li>\n'
                    section_num += 1
        
        if 'discussion' in draft and draft['discussion']:
            html_content += f'            <li><a href="#discussion">{section_num}. Discussion</a></li>\n'
            section_num += 1
        if 'conclusion' in draft and draft['conclusion']:
            html_content += f'            <li><a href="#conclusion">{section_num}. Conclusion</a></li>\n'
        if 'bibliography' in draft and draft['bibliography']:
            html_content += '            <li><a href="#references">References</a></li>\n'
        
        html_content += '        </ul>\n'
        html_content += '    </div>\n\n'
        
        # Abstract with enhanced formatting
        if 'abstract' in draft and draft['abstract']:
            html_content += '    <div class="abstract" id="abstract">\n'
            html_content += '        <h2>Abstract</h2>\n'
            html_content += '        <div class="section-overview">This section provides a concise overview of the research findings and methodology.</div>\n'
            formatted_abstract = self._format_html_content_with_structure(str(draft['abstract']))
            html_content += formatted_abstract
            html_content += '    </div>\n\n'
        
        # Introduction with enhanced formatting
        if 'introduction' in draft and draft['introduction']:
            html_content += '    <h2 id="introduction">1. Introduction</h2>\n'
            html_content += '    <h3>1.1 Background</h3>\n'
            formatted_intro = self._format_html_content_with_structure(str(draft['introduction']))
            html_content += formatted_intro
        
        # Main sections with improved structure
        section_num = 2
        if 'sections' in draft and draft['sections']:
            for key, section in draft['sections'].items():
                if key.startswith('theme_') and isinstance(section, dict):
                    title_clean = section.get('title', f'Section {section_num}')
                    content = section.get('content', 'Content unavailable')
                    
                    html_content += f'    <h2 id="section-{section_num}">{section_num}. {self._escape_html(title_clean)}</h2>\n'
                    html_content += f'    <h3>{section_num}.1 Overview</h3>\n'
                    
                    formatted_content = self._format_html_content_with_structure(str(content))
                    html_content += formatted_content
                    
                    # Add key findings if content is substantial
                    if len(str(content)) > 200:
                        html_content += f'    <h3>{section_num}.2 Key Insights</h3>\n'
                        html_content += '    <div class="key-points">\n'
                        html_content += '        <h4>Key Points:</h4>\n'
                        insights = self._extract_html_key_points(str(content))
                        html_content += insights
                        html_content += '    </div>\n'
                    
                    section_num += 1
        
        # Discussion with enhanced structure
        if 'discussion' in draft and draft['discussion']:
            html_content += f'    <h2 id="discussion">{section_num}. Discussion</h2>\n'
            html_content += f'    <h3>{section_num}.1 Analysis</h3>\n'
            formatted_discussion = self._format_html_content_with_structure(str(draft['discussion']))
            html_content += formatted_discussion
            
            html_content += f'    <h3>{section_num}.2 Implications</h3>\n'
            html_content += '    <div class="implications">\n'
            html_content += '        <p>The findings presented in this research have several important implications:</p>\n'
            html_content += '        <ul>\n'
            html_content += '            <li><span class="strong-point">Theoretical Implications:</span> Further research in this area</li>\n'
            html_content += '            <li><span class="strong-point">Practical Applications:</span> Real-world implementation considerations</li>\n'
            html_content += '            <li><span class="strong-point">Future Directions:</span> Areas for continued investigation</li>\n'
            html_content += '        </ul>\n'
            html_content += '    </div>\n'
            section_num += 1
        
        # Conclusion with enhanced structure
        if 'conclusion' in draft and draft['conclusion']:
            html_content += f'    <h2 id="conclusion">{section_num}. Conclusion</h2>\n'
            html_content += '    <h3>Summary of Findings</h3>\n'
            formatted_conclusion = self._format_html_content_with_structure(str(draft['conclusion']))
            html_content += formatted_conclusion
            
            html_content += '    <h3>Final Remarks</h3>\n'
            html_content += '    <p>This research contributes to the ongoing understanding of the topic and provides a foundation for future investigations.</p>\n'
        
        # References with proper formatting
        if 'bibliography' in draft and draft['bibliography']:
            html_content += '    <div class="references" id="references">\n'
            html_content += '        <h2>References</h2>\n'
            bib_text = str(draft['bibliography'])
            if isinstance(draft['bibliography'], list):
                html_content += '        <ol>\n'
                for ref in draft['bibliography']:
                    html_content += f'            <li>{self._escape_html(str(ref))}</li>\n'
                html_content += '        </ol>\n'
            else:
                html_content += f'        <p>{self._escape_html(bib_text)}</p>\n'
            html_content += '    </div>\n'
        
        # Footer
        html_content += '    <div class="footer">\n'
        html_content += '        <p>Generated by Academic Research Assistant</p>\n'
        html_content += '    </div>\n'
        html_content += '</body>\n</html>'
        
        return html_content
    
    def _format_html_content_with_structure(self, content: str) -> str:
        """Format content with proper HTML paragraph structure"""
        if not content or len(content.strip()) == 0:
            return '        <p>Content not available.</p>\n'
        
        # Split content into sentences
        sentences = [s.strip() for s in content.split('.') if s.strip()]
        
        if len(sentences) <= 2:
            return f'        <p>{self._escape_html(content)}</p>\n'
        
        formatted_content = ""
        
        # Group sentences into paragraphs
        current_paragraph = []
        for i, sentence in enumerate(sentences):
            current_paragraph.append(sentence)
            
            # Create paragraph breaks every 2-3 sentences or at logical breaks
            if (len(current_paragraph) >= 3 or 
                i == len(sentences) - 1 or
                any(keyword in sentence.lower() for keyword in ['however', 'therefore', 'furthermore', 'additionally', 'moreover'])):
                
                paragraph_text = '. '.join(current_paragraph)
                if not paragraph_text.endswith('.'):
                    paragraph_text += '.'
                formatted_content += f'        <p>{self._escape_html(paragraph_text)}</p>\n'
                current_paragraph = []
        
        return formatted_content
    
    def _extract_html_key_points(self, content: str) -> str:
        """Extract key points and format as HTML list items"""
        sentences = [s.strip() for s in content.split('.') if s.strip() and len(s.strip()) > 20]
        
        if len(sentences) < 2:
            return f'        <ul><li>{self._escape_html(content)}</li></ul>\n'
        
        html_content = '        <ul>\n'
        for sentence in sentences[:4]:  # Take up to 4 key sentences
            if len(sentence) > 30:  # Only substantial sentences
                clean_sentence = sentence.strip()
                if not clean_sentence.endswith('.'):
                    clean_sentence += '.'
                html_content += f'            <li><span class="strong-point">{self._escape_html(clean_sentence)}</span></li>\n'
        html_content += '        </ul>\n'
        
        return html_content
    
    def _format_papers_as_bibtex(self, papers: List[Any]) -> str:
        """Format papers as BibTeX entries"""
        bibtex_entries = []
        header = f"""% Bibliography generated by Academic Research Assistant
% Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
% Total entries: {len(papers)}

"""
        
        for i, paper in enumerate(papers):
            try:
                # Generate citation key
                first_author = ""
                if hasattr(paper, 'authors') and paper.authors:
                    first_author = paper.authors[0].split()[-1] if paper.authors[0] else f"author{i}"
                else:
                    first_author = f"author{i}"
                
                year = ""
                if hasattr(paper, 'published_date') and paper.published_date:
                    year = str(paper.published_date.year)
                else:
                    year = "unknown"
                
                citation_key = f"{first_author.lower()}{year}_{i}"
                
                # Create BibTeX entry
                entry_type = "article"
                if hasattr(paper, 'journal') and not paper.journal:
                    entry_type = "misc"
                
                bibtex_entry = f"@{entry_type}{{{citation_key},\n"
                
                # Title
                if hasattr(paper, 'title') and paper.title:
                    bibtex_entry += f"  title={{{self._escape_bibtex(paper.title)}}},\n"
                
                # Authors
                if hasattr(paper, 'authors') and paper.authors:
                    authors = " and ".join(paper.authors)
                    bibtex_entry += f"  author={{{self._escape_bibtex(authors)}}},\n"
                
                # Journal
                if hasattr(paper, 'journal') and paper.journal:
                    bibtex_entry += f"  journal={{{self._escape_bibtex(paper.journal)}}},\n"
                
                # Year
                if year != "unknown":
                    bibtex_entry += f"  year={{{year}}},\n"
                
                # DOI
                if hasattr(paper, 'doi') and paper.doi:
                    bibtex_entry += f"  doi={{{paper.doi}}},\n"
                
                # URL
                if hasattr(paper, 'url') and paper.url:
                    bibtex_entry += f"  url={{{paper.url}}},\n"
                
                # Abstract
                if hasattr(paper, 'abstract') and paper.abstract:
                    abstract = paper.abstract[:500] + "..." if len(paper.abstract) > 500 else paper.abstract
                    bibtex_entry += f"  abstract={{{self._escape_bibtex(abstract)}}},\n"
                
                # Remove trailing comma and close entry
                bibtex_entry = bibtex_entry.rstrip(',\n') + "\n}\n"
                bibtex_entries.append(bibtex_entry)
                
            except Exception as e:
                logger.warning(f"Error creating BibTeX entry for paper {i}: {e}")
                continue
        
        return header + "\n".join(bibtex_entries)
    
    # Utility methods for escaping special characters
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        latex_escapes = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '^': '\\textasciicircum{}',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '\\': '\\textbackslash{}'
        }
        
        result = text
        for char, escape in latex_escapes.items():
            result = result.replace(char, escape)
        return result
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        html_escapes = {
            '&': '&amp;',
            '<': '&lt;',
            '>': '&gt;',
            '"': '&quot;',
            "'": '&#x27;'
        }
        
        result = text
        for char, escape in html_escapes.items():
            result = result.replace(char, escape)
        return result
    
    def _escape_bibtex(self, text: str) -> str:
        """Escape special BibTeX characters"""
        # For BibTeX, we mainly need to handle braces and special characters
        result = text.replace('{', '\\{').replace('}', '\\}')
        return result


# Create global export manager instance
export_manager = ExportManager()
